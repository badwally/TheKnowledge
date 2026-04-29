"""Tests for the Word .docx converter (M30)."""

from __future__ import annotations

from datetime import datetime
from pathlib import Path

import pytest
from docx import Document

from gateway import converters
from gateway import frontmatter as fm
from gateway import paths
from gateway.converters import docx as docx_mod
from gateway.converters.base import ConversionError


@pytest.fixture(autouse=True)
def _reset_registry():
    converters.reset_registry_for_tests()
    yield
    converters.reset_registry_for_tests()


def _build_docx(
    path: Path,
    *,
    paragraphs: list[tuple[str, str | None]] | None = None,
    tables: list[list[list[str]]] | None = None,
    author: str = "",
    title: str = "",
    subject: str = "",
    created: datetime | None = None,
) -> None:
    """Write a .docx file at `path`. paragraphs is [(text, style_or_None)]."""
    doc = Document()
    if paragraphs:
        for text, style in paragraphs:
            if style:
                doc.add_paragraph(text, style=style)
            else:
                doc.add_paragraph(text)
    if tables:
        for grid in tables:
            rows = len(grid)
            cols = len(grid[0]) if grid else 0
            t = doc.add_table(rows=rows, cols=cols)
            for r_idx, row in enumerate(grid):
                for c_idx, cell in enumerate(row):
                    t.rows[r_idx].cells[c_idx].text = cell
    cp = doc.core_properties
    if author:
        cp.author = author
    if title:
        cp.title = title
    if subject:
        cp.subject = subject
    if created is not None:
        cp.created = created
    doc.save(str(path))


# --- detect ----------------------------------------------------------------


def test_docx_detect_accepts_docx_paths():
    c = docx_mod.DocxConverter()
    assert c.detect("/tmp/foo.docx")
    assert c.detect("foo.DOCX")
    assert c.detect("relative/path.docx")


def test_docx_detect_rejects_other_extensions_and_urls():
    c = docx_mod.DocxConverter()
    assert not c.detect("/tmp/foo.doc")  # legacy format not supported
    assert not c.detect("/tmp/foo.pdf")
    assert not c.detect("/tmp/foo.txt")
    assert not c.detect("https://example.com/foo.docx")


# --- convert ---------------------------------------------------------------


def test_docx_convert_extracts_paragraphs_and_writes_sidecar(kb_root, tmp_path):
    src = tmp_path / "doc.docx"
    _build_docx(
        src,
        paragraphs=[
            ("Introduction", "Heading 1"),
            ("Some body text in the introduction.", None),
            ("Background", "Heading 2"),
            ("More body text.", None),
        ],
        author="Andrew Grant",
        title="Architecture Notes",
        created=datetime(2026, 4, 28),
    )

    text = docx_mod.DocxConverter().convert(str(src))
    front, body = fm.parse(text)

    assert front["type"] == "docx"
    assert front["id"].startswith("docx-")
    assert front["title"] == "Architecture Notes"
    assert front["authors"] == ["Andrew Grant"]
    assert front["published_at"] == "2026"
    assert front["meta"]["paragraph_count"] == 4
    assert front["meta"]["table_count"] == 0
    assert front["meta"]["extraction_tool"] == "python-docx"

    assert "# Introduction" in body
    assert "## Background" in body
    assert "Some body text in the introduction." in body

    sidecar = paths.raw_dir_for("docx") / f"{front['id']}.docx"
    assert sidecar.exists()
    assert sidecar.read_bytes() == src.read_bytes()


def test_docx_convert_uses_author_year_title_id_when_metadata_present(kb_root, tmp_path):
    src = tmp_path / "doc.docx"
    _build_docx(
        src,
        paragraphs=[("body", None)],
        author="Smith",
        title="The Big Idea",
        created=datetime(2025, 1, 1),
    )
    text = docx_mod.DocxConverter().convert(str(src))
    front, _ = fm.parse(text)
    assert front["id"].startswith("docx-smith-2025-")
    assert "the-big-idea" in front["id"]


def test_docx_convert_falls_back_to_hash_id_without_metadata(kb_root, tmp_path):
    src = tmp_path / "doc.docx"
    _build_docx(src, paragraphs=[("body content here", None)])
    text = docx_mod.DocxConverter().convert(str(src))
    front, _ = fm.parse(text)
    # No author/title/year → hash fallback
    assert front["id"].startswith("docx-")
    # Hash IDs are exactly 12 hex chars after the prefix
    suffix = front["id"][len("docx-"):]
    assert len(suffix) == 12
    assert all(c in "0123456789abcdef" for c in suffix)


def test_docx_convert_renders_table_as_markdown(kb_root, tmp_path):
    src = tmp_path / "doc.docx"
    _build_docx(
        src,
        paragraphs=[("Before table", None)],
        tables=[[
            ["region", "qty"],
            ["NA", "12"],
            ["EU", "7"],
        ]],
    )
    text = docx_mod.DocxConverter().convert(str(src))
    front, body = fm.parse(text)

    assert front["meta"]["table_count"] == 1
    assert "| region | qty |" in body
    assert "| NA | 12 |" in body


def test_docx_convert_escapes_pipe_in_table_cell(kb_root, tmp_path):
    src = tmp_path / "doc.docx"
    _build_docx(
        src,
        paragraphs=[("intro", None)],
        tables=[[["a", "b"], ["foo|bar", "plain"]]],
    )
    text = docx_mod.DocxConverter().convert(str(src))
    _, body = fm.parse(text)
    assert "foo\\|bar" in body


def test_docx_convert_idempotent_id_for_same_content(kb_root, tmp_path):
    """Two copies of the same docx file produce the same canonical ID."""
    src1 = tmp_path / "first.docx"
    src2 = tmp_path / "second.docx"
    # Same content, no metadata → both fall back to content hash
    _build_docx(src1, paragraphs=[("identical body", None)])
    src2.write_bytes(src1.read_bytes())

    f1, _ = fm.parse(docx_mod.DocxConverter().convert(str(src1)))
    f2, _ = fm.parse(docx_mod.DocxConverter().convert(str(src2)))
    assert f1["id"] == f2["id"]


def test_docx_convert_missing_file_raises(kb_root):
    with pytest.raises(ConversionError):
        docx_mod.DocxConverter().convert("/does/not/exist.docx")


def test_docx_convert_empty_doc_raises(kb_root, tmp_path):
    src = tmp_path / "empty.docx"
    _build_docx(src, paragraphs=[])
    with pytest.raises(ConversionError):
        docx_mod.DocxConverter().convert(str(src))


def test_docx_convert_corrupt_file_raises(kb_root, tmp_path):
    src = tmp_path / "corrupt.docx"
    src.write_bytes(b"not a real docx file")
    with pytest.raises(ConversionError):
        docx_mod.DocxConverter().convert(str(src))


# --- registry --------------------------------------------------------------


def test_dispatch_docx_path_routes_to_docx_converter():
    c = converters.dispatch("/tmp/notes.docx")
    assert isinstance(c, docx_mod.DocxConverter)


def test_dispatch_doc_legacy_format_has_no_converter():
    """Legacy .doc is intentionally unsupported in v1."""
    from gateway.converters import NoConverterError
    with pytest.raises(NoConverterError):
        converters.dispatch("/tmp/old.doc")
